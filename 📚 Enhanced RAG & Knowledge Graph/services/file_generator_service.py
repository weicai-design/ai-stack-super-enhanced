"""
真实的文件生成服务
支持Word/Excel/PPT/PDF等多种格式
"""
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import json


class FileGeneratorService:
    """文件生成服务"""
    
    def __init__(self):
        """初始化文件生成服务"""
        self.output_dir = Path("data/generated_files")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 检查依赖
        self.docx_available = self._check_docx()
        self.xlsx_available = self._check_xlsx()
        self.pdf_available = self._check_pdf()
    
    def _check_docx(self) -> bool:
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def _check_xlsx(self) -> bool:
        try:
            import openpyxl
            return True
        except ImportError:
            return False
    
    def _check_pdf(self) -> bool:
        try:
            from reportlab.pdfgen import canvas
            return True
        except ImportError:
            return False
    
    async def generate_word(
        self,
        content: str,
        title: Optional[str] = None,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        生成Word文档（真实实现）
        
        Args:
            content: 文档内容
            title: 文档标题
            filename: 文件名
            
        Returns:
            生成结果
        """
        if not self.docx_available:
            return {
                "success": False,
                "error": "python-docx未安装",
                "solution": "运行: pip install python-docx"
            }
        
        try:
            from docx import Document
            
            # 创建文档
            doc = Document()
            
            # 添加标题
            if title:
                doc.add_heading(title, 0)
            
            # 添加内容（按段落分割）
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # 保存文件
            if not filename:
                filename = f"document_{int(datetime.now().timestamp())}.docx"
            
            file_path = self.output_dir / filename
            doc.save(str(file_path))
            
            return {
                "success": True,
                "file_path": str(file_path),
                "file_name": filename,
                "file_size": file_path.stat().st_size,
                "file_type": "docx",
                "download_url": f"/api/file/download/{filename}"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def generate_excel(
        self,
        data: List[Dict[str, Any]],
        sheet_name: str = "Sheet1",
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        生成Excel文件（真实实现）
        
        Args:
            data: 数据列表（字典列表）
            sheet_name: 工作表名称
            filename: 文件名
            
        Returns:
            生成结果
        """
        if not self.xlsx_available:
            return {
                "success": False,
                "error": "openpyxl未安装",
                "solution": "运行: pip install openpyxl"
            }
        
        try:
            import pandas as pd
            
            # 创建DataFrame
            df = pd.DataFrame(data)
            
            # 保存文件
            if not filename:
                filename = f"spreadsheet_{int(datetime.now().timestamp())}.xlsx"
            
            file_path = self.output_dir / filename
            df.to_excel(str(file_path), sheet_name=sheet_name, index=False)
            
            return {
                "success": True,
                "file_path": str(file_path),
                "file_name": filename,
                "file_size": file_path.stat().st_size,
                "file_type": "xlsx",
                "rows": len(data),
                "columns": len(data[0]) if data else 0,
                "download_url": f"/api/file/download/{filename}"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def generate_pdf(
        self,
        content: str,
        title: Optional[str] = None,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        生成PDF文件（真实实现）
        
        Args:
            content: 文档内容
            title: 文档标题
            filename: 文件名
            
        Returns:
            生成结果
        """
        if not self.pdf_available:
            return {
                "success": False,
                "error": "reportlab未安装",
                "solution": "运行: pip install reportlab"
            }
        
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # 保存文件
            if not filename:
                filename = f"document_{int(datetime.now().timestamp())}.pdf"
            
            file_path = self.output_dir / filename
            
            # 创建PDF
            c = canvas.Canvas(str(file_path), pagesize=A4)
            width, height = A4
            
            # 设置字体（支持中文）
            try:
                # 尝试使用系统中文字体
                pdfmetrics.registerFont(TTFont('Chinese', '/System/Library/Fonts/STHeiti Light.ttc'))
                c.setFont('Chinese', 12)
            except:
                # 降级到默认字体
                c.setFont('Helvetica', 12)
            
            # 添加标题
            y = height - 50
            if title:
                c.setFont('Chinese' if 'Chinese' in pdfmetrics.getRegisteredFontNames() else 'Helvetica', 16)
                c.drawString(50, y, title)
                y -= 40
                c.setFont('Chinese' if 'Chinese' in pdfmetrics.getRegisteredFontNames() else 'Helvetica', 12)
            
            # 添加内容
            lines = content.split('\n')
            for line in lines:
                if y < 50:  # 换页
                    c.showPage()
                    y = height - 50
                    c.setFont('Chinese' if 'Chinese' in pdfmetrics.getRegisteredFontNames() else 'Helvetica', 12)
                
                # 处理长行（自动换行）
                if len(line) > 80:
                    words = [line[i:i+80] for i in range(0, len(line), 80)]
                    for word in words:
                        c.drawString(50, y, word)
                        y -= 20
                else:
                    c.drawString(50, y, line)
                    y -= 20
            
            c.save()
            
            return {
                "success": True,
                "file_path": str(file_path),
                "file_name": filename,
                "file_size": file_path.stat().st_size,
                "file_type": "pdf",
                "download_url": f"/api/file/download/{filename}"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def generate_markdown(
        self,
        content: str,
        title: Optional[str] = None,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """生成Markdown文件"""
        try:
            if not filename:
                filename = f"document_{int(datetime.now().timestamp())}.md"
            
            file_path = self.output_dir / filename
            
            # 写入内容
            with open(file_path, 'w', encoding='utf-8') as f:
                if title:
                    f.write(f"# {title}\n\n")
                f.write(content)
            
            return {
                "success": True,
                "file_path": str(file_path),
                "file_name": filename,
                "file_size": file_path.stat().st_size,
                "file_type": "markdown",
                "download_url": f"/api/file/download/{filename}"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def get_status(self) -> Dict[str, Any]:
        """获取文件生成服务状态"""
        return {
            "word_available": self.docx_available,
            "excel_available": self.xlsx_available,
            "pdf_available": self.pdf_available,
            "markdown_available": True,
            "supported_formats": [
                fmt for fmt, available in [
                    ("docx", self.docx_available),
                    ("xlsx", self.xlsx_available),
                    ("pdf", self.pdf_available),
                    ("md", True)
                ]
                if available
            ],
            "installation_guide": {
                "word": "pip install python-docx",
                "excel": "pip install openpyxl pandas",
                "pdf": "pip install reportlab"
            }
        }


# 全局文件生成服务实例
_file_generator = None

def get_file_generator() -> FileGeneratorService:
    """获取文件生成服务实例"""
    global _file_generator
    if _file_generator is None:
        _file_generator = FileGeneratorService()
    return _file_generator


# 使用示例
if __name__ == "__main__":
    import asyncio
    
    async def test():
        generator = get_file_generator()
        
        print("✅ 文件生成服务已加载")
        print(f"📊 状态: {generator.get_status()}")
        
        # 测试生成Word
        if generator.docx_available:
            result = await generator.generate_word(
                content="这是测试内容\n\n第二段内容",
                title="测试文档"
            )
            if result["success"]:
                print(f"\n✅ Word生成成功: {result['file_path']}")
            else:
                print(f"\n❌ Word生成失败: {result['error']}")
        
        # 测试生成Markdown
        result = await generator.generate_markdown(
            content="# 测试\n\n这是内容",
            title="测试文档"
        )
        if result["success"]:
            print(f"\n✅ Markdown生成成功: {result['file_path']}")
    
    asyncio.run(test())


